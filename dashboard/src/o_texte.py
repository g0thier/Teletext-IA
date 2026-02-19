import io
import os

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def transcription(audio_converted, filename):
    base_name = os.path.splitext(filename)[0]
    txt_filename = f"{base_name}_transcription.txt"

    # Création du fichier en mémoire
    txt_bytes = io.BytesIO()

    # Écriture "comme dans le fichier txt", mais dans le buffer mémoire
    for segment in audio_converted:
        parole = segment["text"]
        txt_bytes.write((parole + "\n").encode("utf-8"))

    txt_bytes.seek(0)

    return txt_filename, txt_bytes

def description(video_converted, filename):
    base_name = os.path.splitext(filename)[0]
    txt_filename = f"{base_name}_description.txt"
    out = []

    for segment in video_converted:
        out.append(f"Scene {int(segment["id"])} from {segment["start"]} to {segment["end"]}")
        out.append(f"↳ {segment["text"].capitalize()}.")
        out.append("")  # ligne vide entre blocs

    texte = "\n".join(out)

    # Création du fichier en mémoire
    txt_bytes = io.BytesIO()
    txt_bytes.write(texte.encode("utf-8"))
    txt_bytes.seek(0)

    return txt_filename, txt_bytes

## 
## Fonctions pour full_accessibility()
##

def seconds_to_timecode(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = seconds % 60
    return f"{hours:02}:{minutes:02}:{secs:06.3f}"

def create_styles(document):
    styles = document.styles
    police = "Courier New"

    # =========================
    # VIDEO TIMESTAMP STYLE
    # =========================
    style = styles.add_style("VIDEO_TIMESTAMP", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = police
    style.font.size = Pt(8)
    style.paragraph_format.space_after = Pt(6)

    # =========================
    # VIDEO TITLE STYLE
    # =========================
    style = styles.add_style("VIDEO_TITLE", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = police
    style.font.size = Pt(12)
    style.font
    style.font.bold = True

    # =========================
    # CHARACTER PLACEHOLDER
    # =========================
    style = styles.add_style("CHARACTER_PLACEHOLDER", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = police
    style.font.size = Pt(12)
    style.font.bold = True
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================
    # AUDIO DIALOGUE STYLE
    # =========================
    style = styles.add_style("AUDIO_DIALOGUE", WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = police
    style.font.size = Pt(12)
    style.paragraph_format.left_indent = Cm(2)
    style.paragraph_format.right_indent = Cm(2)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_cinema_margins(document):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)


def export_scenario(json_data, docx_bytes):

    document = Document()

    # Marges cinéma standard
    set_cinema_margins(document)

    # Création des styles personnalisés
    create_styles(document)

    video_counter = 1

    for index, item in enumerate(json_data):
        # Bordure
        next_item = json_data[index + 1]["type"] if index + 1 < len(json_data) else None

        # =========================
        # VIDEO BLOCK
        # =========================
        if item["type"] == "video" and next_item != "video":

            # Numérotation automatique
            timestamp_text = f"{video_counter:03d}  {item['start']} - {item['end']}"
            video_counter += 1

            # Timestamp
            p_time = document.add_paragraph(timestamp_text, style="VIDEO_TIMESTAMP")
            p_time.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Title
            p_text = document.add_paragraph(item["text"], style="VIDEO_TITLE")
            p_text.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # =========================
        # AUDIO BLOCK
        # =========================
        elif item["type"] == "audio":

            # Placeholder personnage
            document.add_paragraph("...", style="CHARACTER_PLACEHOLDER")

            # Dialogue
            document.add_paragraph(item["text"], style="AUDIO_DIALOGUE")

            # Space
            document.add_paragraph("")

    document.save(docx_bytes)

## 
## Fonctions Accessible 
##

def script(video_converted, audio_converted, filename):
    base_name = os.path.splitext(filename)[0]
    txt_filename = f"{base_name}_script.docx"

    audio_highlighted = [
        {
            "type": "audio",
            "start": seconds_to_timecode(e.get("start")),
            "end": seconds_to_timecode(e.get("end")),
            "text": e.get("text").strip(),
        }
        for e in audio_converted
    ]

    video_highlighted = [
        {
            "type": "video",
            "start": e.get("start").get_timecode(),
            "end": e.get("end").get_timecode(),
            "text": e.get("text").upper(),
        }
        for e in video_converted
    ]

    media_highlighted = video_highlighted + audio_highlighted
    media_highlighted.sort(key=lambda x: x["start"])


    # Création du fichier en mémoire
    docx_bytes = io.BytesIO()
    docx_bytes.seek(0)

    # Ecriture
    export_scenario(media_highlighted, docx_bytes)
    docx_bytes.seek(0)

    return txt_filename, docx_bytes
